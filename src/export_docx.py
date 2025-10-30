"""
DOCX Exporter for AirPilot_AI
Converts generated markdown/plaintext drafts into formatted Word files.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def markdown_to_docx(md_text: str, out_path: Path) -> Path:
    """
    Very light Markdown → DOCX converter.
    Supports headings (#, ##, ###), bold **text**, italic *text*,
    unordered lists (- or *), and normal paragraphs.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    for line in lines:
        if not line.strip():
            doc.add_paragraph("")  # blank line
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("-", "*")):
            # unordered list item
            text = line.lstrip("-* ").strip()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
        elif line.startswith("**") and line.endswith("**"):
            run = doc.add_paragraph().add_run(line.strip("*"))
            run.bold = True
        else:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"✓ DOCX written to: {out_path}")
    return out_path


def convert_draft_md_to_docx(md_path: Path) -> Path:
    """Reads an existing .md draft and exports it to DOCX."""
    text = Path(md_path).read_text(encoding="utf-8", errors="ignore")
    out = md_path.with_suffix(".docx")
    return markdown_to_docx(text, out)
